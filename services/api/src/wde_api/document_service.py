"""Trusted deterministic renderers for Phase 10 canonical document data; never fetches or executes source content."""

from __future__ import annotations

import html
import io
import json
from typing import Any

from docx import Document
from docx.shared import Inches
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from wde_api.document_types import CanonicalDocument, DocumentProfile
from wde_api.export_errors import ExportTooLarge, ExportUnsupportedFormat
from wde_api.export_types import CanonicalExportDataset


def _text(value: Any, null_text: str) -> str:
    if value is None:
        return null_text
    if isinstance(value, (dict, list)):
        return json.dumps(value, ensure_ascii=False, sort_keys=True, separators=(",", ":"))
    return str(value).replace("\x00", "")


def build_document(
    dataset: CanonicalExportDataset, *, profile: DocumentProfile = DocumentProfile()
) -> CanonicalDocument:
    return CanonicalDocument(
        profile,
        dataset,
        {
            "schema_version": "document.v1",
            "record_count": len(dataset.rows),
            "validation_run_id": dataset.validation_run_id,
            "plan_version": dataset.plan_version,
        },
    )


def _markdown(document: CanonicalDocument) -> str:
    fields = document.dataset.fields
    lines = [
        f"# {document.profile.title}",
        "",
        "## Summary",
        f"- Records: {len(document.dataset.rows)}",
        f"- Validation run: {document.dataset.validation_run_id}",
        "",
        "## Records",
    ]
    if not document.dataset.rows:
        return "\n".join(lines + ["", "No records found.", ""])
    lines += ["", "| " + " | ".join(fields) + " |", "|" + "|".join("---" for _ in fields) + "|"]
    for row in document.dataset.rows:
        lines.append(
            "| "
            + " | ".join(
                _text(row.get(field), document.profile.null_text).replace("|", "\\|").replace("\n", "<br>")
                for field in fields
            )
            + " |"
        )
    return "\n".join(lines) + "\n"


def render_document(document: CanonicalDocument, format_name: str, *, max_records: int = 10_000) -> bytes:
    if len(document.dataset.rows) > max_records:
        raise ExportTooLarge("Document dataset exceeds the configured record limit.")
    if format_name == "md":
        return _markdown(document).encode("utf-8")
    if format_name == "txt":
        return _markdown(document).replace("# ", "").replace("## ", "").replace("|", " ").encode("utf-8")
    if format_name == "html":
        fields = document.dataset.fields
        rows = "".join(
            "<tr>"
            + "".join(
                f"<td>{html.escape(_text(row.get(field), document.profile.null_text))}</td>"
                for field in fields
            )
            + "</tr>"
            for row in document.dataset.rows
        )
        headers = "".join(f"<th>{html.escape(field)}</th>" for field in fields)
        return f"<!doctype html><html><head><meta charset='utf-8'><title>{html.escape(document.profile.title)}</title></head><body><h1>{html.escape(document.profile.title)}</h1><p>Records: {len(document.dataset.rows)}</p><table><thead><tr>{headers}</tr></thead><tbody>{rows}</tbody></table></body></html>".encode()
    if format_name == "docx":
        return _docx(document)
    if format_name == "pdf":
        return _pdf(document)
    raise ExportUnsupportedFormat("Only pdf, docx, md, txt, and html document formats are supported.")


def _docx(document: CanonicalDocument) -> bytes:
    file = io.BytesIO()
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    doc.add_heading(document.profile.title, level=0)
    doc.add_paragraph(
        f"Records: {len(document.dataset.rows)}  •  Validation run: {document.dataset.validation_run_id}"
    )
    if not document.dataset.rows:
        doc.add_paragraph("No records found.")
    else:
        table = doc.add_table(rows=1, cols=len(document.dataset.fields))
        table.style = "Table Grid"
        for cell, field in zip(table.rows[0].cells, document.dataset.fields, strict=True):
            cell.text = field
        for row in document.dataset.rows:
            cells = table.add_row().cells
            for cell, field in zip(cells, document.dataset.fields, strict=True):
                cell.text = _text(row.get(field), document.profile.null_text)
    doc.save(file)
    return file.getvalue()


def _pdf(document: CanonicalDocument) -> bytes:
    output = io.BytesIO()
    wide = len(document.dataset.fields) > 5
    page = landscape(A4) if wide else A4
    styles = getSampleStyleSheet()
    story = [
        Paragraph(document.profile.title, styles["Title"]),
        Paragraph(
            f"Records: {len(document.dataset.rows)} | Validation run: {document.dataset.validation_run_id}",
            styles["BodyText"],
        ),
        Spacer(1, 0.35 * cm),
    ]
    if not document.dataset.rows:
        story.append(Paragraph("No records found.", styles["BodyText"]))
    else:
        data = [[Paragraph(html.escape(field), styles["BodyText"]) for field in document.dataset.fields]]
        for row in document.dataset.rows:
            data.append(
                [
                    Paragraph(
                        html.escape(_text(row.get(field), document.profile.null_text)), styles["BodyText"]
                    )
                    for field in document.dataset.fields
                ]
            )
        table = Table(
            data,
            repeatRows=1,
            colWidths=[(page[0] - 3 * cm) / len(document.dataset.fields)] * len(document.dataset.fields),
        )
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e7e5df")),
                    ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#b9b6af")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("FONTSIZE", (0, 0), (-1, -1), 7),
                ]
            )
        )
        story.append(table)
    SimpleDocTemplate(
        output,
        pagesize=page,
        leftMargin=1.5 * cm,
        rightMargin=1.5 * cm,
        topMargin=1.5 * cm,
        bottomMargin=1.5 * cm,
    ).build(story)
    return output.getvalue()
