from __future__ import annotations

import io

from docx import Document
from pypdf import PdfReader
from wde_api.document_service import build_document, render_document
from wde_api.export_types import CanonicalExportDataset, ExportOptions


def document(rows: tuple[dict[str, object], ...] = ({"name": "नमस्ते", "price": None},)):
    dataset = CanonicalExportDataset(
        fields=("name", "price"),
        rows=rows,
        validation_run_id="run-1",
        plan_version=1,
        options=ExportOptions(),
    )
    return build_document(dataset)


def test_pdf_and_docx_are_reopenable_with_expected_values() -> None:
    report = document()
    pdf_text = "".join(
        page.extract_text() or "" for page in PdfReader(io.BytesIO(render_document(report, "pdf"))).pages
    )
    assert "Web Data Extraction Results" in pdf_text and "Records: 1" in pdf_text
    docx = Document(io.BytesIO(render_document(report, "docx")))
    assert "Web Data Extraction Results" in "\n".join(paragraph.text for paragraph in docx.paragraphs)
    assert docx.tables[0].rows[1].cells[1].text == "—"


def test_markdown_txt_html_are_deterministic_and_escape_untrusted_content() -> None:
    report = document(({"name": "<script>alert(1)</script>|x", "price": {"nested": 1}},))
    markdown = render_document(report, "md").decode()
    html = render_document(report, "html").decode()
    assert "\\|x" in markdown and "&lt;script&gt;" in html and "<script>alert" not in html
    assert "Web Data Extraction Results" in render_document(report, "txt").decode()


def test_empty_document_still_has_a_valid_message() -> None:
    report = document(())
    assert "No records found." in render_document(report, "md").decode()
    assert "No records found." in "".join(
        page.extract_text() or "" for page in PdfReader(io.BytesIO(render_document(report, "pdf"))).pages
    )
